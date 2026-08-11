import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_docx():
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Fonts
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Plus Jakarta Sans'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22) # Charcoal

    # Title Box (Table with 1 cell and border)
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = title_table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_background(cell, "1E1E1E") # Dark background
    set_cell_margins(cell, top=200, bottom=200, left=200, right=200)

    # Border for cell
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="18" w:space="0" w:color="D4AF37"/><w:left w:val="single" w:sz="18" w:space="0" w:color="D4AF37"/><w:bottom w:val="single" w:sz="18" w:space="0" w:color="D4AF37"/><w:right w:val="single" w:sz="18" w:space="0" w:color="D4AF37"/></w:tcBorders>')
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run("PROPOSAL PEMBUATAN WEB\n")
    run1.bold = True
    run1.font.size = Pt(16)
    run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    run2 = p.add_run('"Kedai Marjuki\'S - E-Commerce Kuliner Rumahan"\n\n')
    run2.bold = True
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37) # Gold

    run3 = p.add_run("Oleh: Arjunaa / Tim Pengembang Kedai Marjuki'S")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_heading(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C) # Dark Red / Gold Accent
        return h

    def add_subheading(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.bold = True
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37) # Gold
        return h

    def add_body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        r2 = p.add_run(text)
        return p

    # PENDAHULUSAN
    add_heading("PENDAHULUSAN")
    add_body("Kedai Marjuki'S (E-Commerce Kuliner Rumahan) merupakan aplikasi sistem penjualan dan pemesanan makanan berbasis web modern yang dikembangkan untuk menjadi solusi digital pintar bagi unit usaha kuliner rumahan Kedai Marjuki'S yang berlokasi di Candi, Semarang. Aplikasi ini tidak hanya berfungsi sebagai katalog online produk hidangan favorit, tetapi juga sebagai platform manajemen operasional komprehensif yang menghadirkan pengalaman efisien dalam mengelola keranjang belanja, proses checkout, verifikasi pembayaran digital (QRIS), hingga rekapitulasi penjualan secara otomatis.")
    add_body("Melalui kombinasi teknologi Full-Stack (Laravel Native, PHP, MySQL Database, HTML5, Vanilla CSS3, JavaScript, dan Bootstrap 5) serta antarmuka visual yang estetik dan modern, Kedai Marjuki'S menghadirkan layanan pemesanan makanan yang responsif, presisi, mudah diakses, dan relevan dengan kebutuhan pelanggan kuliner masa kini.")
    add_body('Dengan mengusung tagline "Cita Rasa Halaman Rumah", aplikasi ini bertujuan membantu pelanggan dalam melakukan pemesanan makanan dan minuman secara daring serta memudahkan pemilik kedai (Admin) dalam mempercepat verifikasi pembayaran, mengontrol stok produk, memantau riwayat transaksi harian, dan memastikan transparansi penjualan dalam satu layar dashboard.')

    # PROFIL PERUSAHAAN
    add_heading("PROFIL PERUSAHAAN")
    add_bullet("Nama Usaha / Perusahaan: ", "Kedai Marjuki'S Digital")
    add_bullet("Bidang Usaha: ", "Culinary E-Commerce & Web Application Development")
    add_bullet("Alamat Kedai: ", "JL. Jomblang Perbalan No 800 Candi, Candisari, Kota Semarang, Jawa Tengah")
    add_bullet("Lokasi Usaha: ", "Berada tepat di halaman rumah")
    add_bullet("Kontak WhatsApp: ", "0882005116301")
    add_bullet("Email: ", "info@marjukis.test / marjukis@gmail.com")
    add_bullet("Website: ", "www.kedai-marjukis.test (Localhost: http://localhost:8000)")
    add_bullet("Portofolio: ", "Sebelumnya, tim pengembang telah berhasil mengembangkan berbagai sistem e-commerce dan aplikasi manajemen inventaris berbasis web yang mampu menangani transaksi harian secara cepat, aman, dan terintegrasi basis data relasional.")
    add_body("Dengan pengalaman tersebut, tim pengembang terus berinovasi menghadirkan Kedai Marjuki'S sebagai platform e-commerce kuliner rumahan yang interaktif, informatif, presisi, dan mudah digunakan baik oleh pembeli maupun pemilik kedai.")

    # DESKRIPSI APLIKASI
    add_heading("DESKRIPSI APLIKASI")
    add_bullet("Nama Aplikasi: ", "Kedai Marjuki'S (E-Commerce Kuliner Rumahan)")
    add_bullet("Jenis Aplikasi: ", "Aplikasi E-Commerce Web, Sistem Kasir (POS) & Manajemen Pesanan Kuliner.")
    
    add_subheading("Fitur Utama:")
    add_bullet("1. Katalog Produk & Filter Kategori Estetik – ", "Memfasilitasi pembeli dalam menjelajahi berbagai menu makanan favorit (Soto Ayam, Nasi Rames, Indomie Biasa/Telor) dan minuman segar (Es Teh, Es Jeruk) dengan filter kategori cepat serta pencarian produk interaktif.")
    add_bullet("2. Keranjang Belanja & Checkout Realtime – ", "Mendukung penambahan produk ke keranjang belanja dengan kalkulasi total harga otomatis, manajemen jumlah porsi (quantity), serta pengisian data pemesan yang cepat.")
    add_bullet("3. Sistem Pembayaran Ganda (QRIS Digital & Cash/Tunai) – ", "Menyediakan opsi pembayaran instan melalui scan QRIS (GoPay, OVO, ShopeePay, Dana, Bank) dengan unggah bukti transfer serta opsi bayar tunai langsung saat pesanan diterima.")
    add_bullet("4. Panel Admin & Verifikasi Pembayaran Kilat – ", "Memungkinkan pemilik kedai untuk mengonfirmasi atau menolak bukti pembayaran QRIS, memperbarui status pesanan (Diproses, Siap, Selesai), mengelola data produk (CRUD), serta mengatur gambar QRIS aktif.")
    add_bullet("5. Dashboard Visual & Laporan Penjualan – ", "Menyajikan rekapitulasi total pendapatan kotor, jumlah pesanan masuk, total produk terdaftar, serta riwayat transaksi harian yang transparan.")

    add_subheading("Keamanan Data:")
    add_body("Kedai Marjuki'S sangat memperhatikan aspek keamanan dan kerahasiaan data pengguna dengan menerapkan sistem perlindungan data berlapis, antara lain:")
    add_bullet("1. Enkripsi Password Hash: ", "Menggunakan enkripsi Bcrypt / SHA-256 untuk keamanan akun Admin & Pembeli.")
    add_bullet("2. Pembatasan Hak Akses (Role-Based Access Control): ", "Akun Pembeli hanya diperbolehkan mengakses menu belanja dan pesanan pribadi, sementara Admin memiliki akses penuh atas manajemen kedai.")
    add_bullet("3. Integritas Basis Data Relasional: ", "Menyimpan data transaksi dan pelanggan secara terenkripsi pada Database MySQL Relasional yang terjamin kestabilannya.")
    add_bullet("4. Proteksi CSRF Token: ", "Menerapkan proteksi keamanan form pada setiap transaksi untuk mencegah manipulasi data dari pihak luar.")

    add_subheading("Manfaat untuk Pengguna:")
    add_bullet("1. ", "Membantu pemilik kedai memantau transaksi dan pendapatan secara realtime.")
    add_bullet("2. ", "Mempercepat proses pemesanan makanan tanpa antre di meja kedai.")
    add_bullet("3. ", "Mengeliminasi kesalahan perhitungan harga dan kembalian secara otomatis.")
    add_bullet("4. ", "Menyediakan kemudahan pembayaran digital via QRIS untuk pelanggan modern.")
    add_bullet("5. ", "Memudahkan rekapitulasi pembukuan dan stok menu harian.")

    # KEUNGGULAN PRODUK
    add_heading("KEUNGGULAN PRODUK")
    add_body("Kedai Marjuki'S memiliki berbagai keunggulan yang menjadikannya berbeda dari aplikasi e-commerce umum lainnya, baik dari sisi teknologi, tampilan visual, maupun pengalaman pengguna.")

    add_subheading("A. Keunggulan Fungsional")
    add_bullet("• Kalkulasi Otomatis Presisi: ", "Perhitungan total harga menu, rincian item belanja, dan total tagihan pesanan dihitung secara realtime.")
    add_bullet("• Pembaruan Status Transparan: ", "Pelanggan dapat memantau status pembayaran dan status memasak pesanan secara langsung dari dashboard pembeli.")
    add_bullet("• Service Layer OOP Architecture: ", "Pemisahan logika bisnis (CartService, OrderService, PaymentService) yang rapi sehingga transaksi berjalan stabil dan bebas kesalahan.")

    add_subheading("B. Keunggulan Teknis")
    add_bullet("• Teknologi Responsif & Fast Load: ", "Dibangun menggunakan Laravel Native dan Bootstrap 5 yang menghasilkan kecepatan muat halaman sangat cepat di berbagai perangkat (HP, Tablet, Desktop).")
    add_bullet("• Arsitektur Database Relasional Handal: ", "Terhubung ke MySQL Database dengan foreign key constraints yang menjaga konsistensi data pesanan.")
    add_bullet("• Struktur Kode Modular: ", "Penggunaan komponen Blade (Blade Components & Layouts) yang mudah dikembangkan dan dirawat jangka panjang.")

    add_subheading("C. Keunggulan Desain & Pengalaman Pengguna (UI/UX)")
    add_bullet("• Antarmuka Dark & Warm Aesthetic: ", "Kombinasi warna gelap elegan (Dark Slate) dengan aksen warna hangat Merah & Emas (#dc2626 & #f59e0b) memberikan kesan modern, bersih, dan menggugah selera.")
    add_bullet("• Integrasi Foto Asli Kedai & Etalase: ", "Menampilkan foto asli warung rumahan dan etalase hidangan Kedai Marjuki'S untuk membangun kepercayaan pelanggan.")
    add_bullet("• Desain Glassmorphic Modern: ", "Penggunaan efek backdrop-blur dan card shadow yang halus untuk pengalaman visual tingkat tinggi.")

    add_subheading("D. Keunggulan Strategis")
    add_bullet("• Mendukung Digitalisasi UMKM Kuliner Rumahan: ", "Memberikan solusi digitalisasi toko makanan berbiaya efisien bagi usaha kuliner lokal.")
    add_bullet("• Transaksi Langsung Tanpa Potongan Aplikasi: ", "Pemilik kedai menerima pembayaran 100% penuh tanpa dipotong biaya komisi pihak ketiga.")

    # SKEMA PEMASARAN
    add_heading("SKEMA PEMASARAN")
    add_bullet("1. Promosi Melalui Media Sosial & Portofolio Digital – ", "Memanfaatkan platform Instagram, TikTok, dan WhatsApp Story untuk memamerkan foto dan video sajian hangat Kedai Marjuki'S yang terhubung langsung ke link website pemesanan.")
    add_bullet("2. Spanduk & Banner QR Code di Halaman Rumah Kedai – ", "Memasang banner fisik ber-QR Code di area Kedai Marjuki'S (JL. Jomblang Perbalan No 800) sehingga pengunjung dapat langsung me-scan dan memesan melalui smartphone.")
    add_bullet("3. Program Penawaran Khusus Pelanggan Rumahan – ", "Menyediakan promo diskon atau paket hemat menu favorit bagi pelanggan sekitar wilayah Candi dan Candisari Semarang yang memesan secara online.")
    add_bullet("4. Edukasi Pemesanan Online Praktis – ", "Menyajikan petunjuk pemesanan dan pembayaran QRIS yang simpel sehingga dapat digunakan dengan mudah oleh semua kalangan masyarakat.")

    # LAYANAN PURNA JUAL
    add_heading("LAYANAN PURNA JUAL")
    add_body("Sebagai bentuk komitmen terhadap kepuasan dan kepercayaan pengguna, Kedai Marjuki'S Tech menyediakan layanan purna jual yang komprehensif. Layanan ini bertujuan untuk memastikan aplikasi berjalan optimal, mudah digunakan, dan terus berkembang sesuai kebutuhan kedai.")

    add_subheading("1. Garansi Bug Fixing")
    add_body("Kami memberikan garansi perbaikan kesalahan sistem (bug) tanpa biaya tambahan selama masa penggunaan. Setiap laporan bug akan ditangani maksimal 2x24 jam oleh tim teknis kami untuk menjaga kelancaran transaksi kedai.")

    add_subheading("2. Maintenance & Update Berkala")
    add_body("Tim pemeliharaan melakukan pemeliharaan sistem (maintenance) dan pembaruan fitur (update) secara rutin untuk menjaga performa dan keamanan basis data MySQL.")

    add_subheading("3. Training Penggunaan Aplikasi")
    add_body("Kami menyediakan sesi pelatihan (training) bagi pemilik kedai dan admin terkait cara penggunaan aplikasi, pengelolaan data menu & kategori, verifikasi pembayaran QRIS, serta pencetakan laporan.")
    add_bullet("• ", "Dapat dilakukan secara online (Zoom/Google Meet) maupun tatap muka langsung di kedai.")
    add_bullet("• ", "Disertai buku panduan operasional (user manual) lengkap.")

    add_subheading("4. Customer Support")
    add_body("Tim dukungan pelanggan siap membantu setiap kendala teknis maupun administratif melalui berbagai saluran komunikasi:")

    # Table for Customer Support
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Media Layanan", "Waktu Layanan", "Keterangan"]
    data = [
        ["WhatsApp Support", "08.00 – 21.00 WIB", "Fast response untuk kendala teknis operasional kasir/web"],
        ["Email Support", "24 Jam", "Penanganan resmi untuk laporan bug dan update database"],
        ["Telepon Langsung", "Jam Kerja (Senin – Minggu)", "Konsultasi langsung dengan tim teknis pengembang"]
    ]

    # Style Header Row
    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        set_cell_background(hdr_cells[idx], "D4AF37") # Gold Header
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Style Data Rows
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        bg_hex = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_hex)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_subheading("5. Komitmen Layanan")
    add_body("Kami berkomitmen memberikan layanan yang cepat, ramah, dan solutif. Tim pengembang berusaha memastikan pemilik kedai dan pelanggan mendapatkan pengalaman terbaik — mulai dari tahap implementasi awal hingga penggunaan harian.")

    # PENUTUP
    add_heading("PENUTUP")
    add_body("Melalui aplikasi Kedai Marjuki'S (E-Commerce Kuliner Rumahan), kami berkomitmen menghadirkan solusi digital pemesanan makanan yang inovatif, estetik, dan efisien di bidang usaha kuliner rumahan. Aplikasi ini tidak hanya membantu pelanggan dalam menikmati sajian lezat dengan mudah, tetapi juga menjadi alat bantu utama bagi pemilik usaha dalam menganalisis dan mengembangkan bisnis secara terukur.")
    add_body("Kami percaya bahwa kehadiran aplikasi ini dapat memberikan nilai tambah yang signifikan bagi efisiensi operasional dan transparansi keuangan Kedai Marjuki'S. Dengan dukungan teknologi Full-Stack yang modern, antarmuka premium, serta layanan purna jual yang responsif, kami siap berkolaborasi untuk menciptakan ekosistem e-commerce kuliner rumahan yang aman, nyaman, dan bermanfaat.")
    add_body("Kami berharap proposal ini dapat menjadi landasan kerja sama dan penilaian yang baik dalam rangka Uji Kompetensi Keahlian (UKK) PPLG 2026. Terima kasih atas perhatian dan kesempatan yang diberikan.")

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(20)
    r_sig1 = p_sig.add_run("Hormat kami,\n\n\n\n")
    r_sig2 = p_sig.add_run("Arjunaa - Tim Pengembang Kedai Marjuki'S")
    r_sig2.bold = True
    r_sig2.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)

    output_path = "Proposal_Pembuatan_Web_Kedai_Marjukis.docx"
    doc.save(output_path)
    print("DOCX_CREATED_SUCCESS: " + output_path)

if __name__ == '__main__':
    create_docx()
