import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

doc = docx.Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Color Palette
PRIMARY_COLOR = RGBColor(185, 28, 28)     # Crimson Red
SECONDARY_COLOR = RGBColor(30, 41, 59)   # Slate Dark
TEXT_COLOR = RGBColor(51, 65, 85)        # Text Slate
MUTED_COLOR = RGBColor(100, 116, 139)    # Muted Gray

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_header_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Plus Jakarta Sans'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    return p

def add_header_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Plus Jakarta Sans'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR
    return p

def add_body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Plus Jakarta Sans'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = SECONDARY_COLOR
    run = p.add_run(text)
    run.font.name = 'Plus Jakarta Sans'
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_COLOR
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Plus Jakarta Sans'
        r_pre.font.size = Pt(10)
        r_pre.font.bold = True
        r_pre.font.color.rgb = SECONDARY_COLOR
    run = p.add_run(text)
    run.font.name = 'Plus Jakarta Sans'
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_COLOR
    return p

# ----------------- COVER / TITLE -----------------
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(30)
title_p.paragraph_format.space_after = Pt(4)
r_title = title_p.add_run("PROPOSAL PENGEMBANGAN SISTEM INFORMASI")
r_title.font.name = 'Plus Jakarta Sans'
r_title.font.size = Pt(18)
r_title.font.bold = True
r_title.font.color.rgb = PRIMARY_COLOR

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(20)
r_sub = sub_p.add_run("E-COMMERCE & DINE-IN ORDERING SYSTEM\nKEDAI MARJUKI'S BERBASIS WEB (LARAVEL 11)")
r_sub.font.name = 'Plus Jakarta Sans'
r_sub.font.size = Pt(14)
r_sub.font.bold = True
r_sub.font.color.rgb = SECONDARY_COLOR

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_p.paragraph_format.space_after = Pt(30)
r_auth = author_p.add_run("Disusun Oleh:\nARJUNA CHYAH\nProgram Keahlian Pengembangan Perangkat Lunak & Gim (PPLG)\n2026")
r_auth.font.name = 'Plus Jakarta Sans'
r_auth.font.size = Pt(11)
r_auth.font.bold = True
r_auth.font.color.rgb = MUTED_COLOR

doc.add_page_break()

# ----------------- DAFTAR ISI -----------------
add_header_1("DAFTAR ISI")
toc_items = [
    "BAB I: PENDAHULUAN",
    "  1.1 Latar Belakang Masalah",
    "  1.2 Rumusan Masalah",
    "  1.3 Maksud dan Tujuan",
    "  1.4 Manfaat Sistem (Pelanggan, Kasir, Admin, Owner)",
    "  1.5 Ruang Lingkup & Batasan Masalah",
    "BAB II: LANDASAN TEORI & SPESIFIKASI TEKNIS",
    "  2.1 Landasan Teori Teknologi (PHP, Laravel 11, SQLite, Bootstrap 5.3, QRIS GPN)",
    "  2.2 Spesifikasi Kebutuhan Perangkat Keras (Hardware)",
    "  2.3 Spesifikasi Kebutuhan Perangkat Lunak (Software)",
    "  2.4 Matriks Hak Akses Pengguna (User Roles Matrix)",
    "BAB III: PERANCANGAN SISTEM (SYSTEM DESIGN)",
    "  3.1 Context Diagram (Diagram Konteks Level 0)",
    "  3.2 Data Flow Diagram (DFD Level 1)",
    "  3.3 Use Case Diagram & Skenario Aktor",
    "  3.4 Conceptual Data Model (CDM)",
    "  3.5 Physical Data Model (PDM)",
    "  3.6 Entity Relationship Diagram (ERD Notasi Chen)",
    "  3.7 Flowchart Workflow Operasional & Shift",
    "BAB IV: RENCANA ANGGARAN BIAYA (RAB) & JADWAL PELAKSANAAN",
    "  4.1 Rencana Anggaran Biaya (RAB Pengembangan & Maintenance)",
    "  4.2 Timeline Pelaksanaan Proyek (Gantt Chart 8 Minggu)",
    "BAB V: PENGUJIAN SISTEM (BLACK BOX TESTING)",
    "  5.1 Metode Pengujian Sistem",
    "  5.2 Matriks Skenario Test Case Black Box",
    "BAB VI: PENUTUP",
    "  6.1 Kesimpulan",
    "  6.2 Saran & Rencana Pengembangan Masa Depan"
]
for item in toc_items:
    add_body(item)

doc.add_page_break()

# ----------------- BAB I -----------------
add_header_1("BAB I: PENDAHULUAN")

add_header_2("1.1 Latar Belakang Masalah")
add_body("Sektor Usaha Mikro, Kecil, dan Menengah (UMKM) bidang kuliner rumahan memegang peranan krusial dalam perekonomian lokal. Kedai Marjuki'S merupakan unit usaha kuliner rumahan yang beralamat di JL. Jomblang Perbalan No 800 Candi, Candisari, Semarang, Jawa Tengah. Kedai ini menyajikan aneka hidangan sarapan dan makan santai khas keluarga seperti Soto Ayam, Nasi Rames, Indomie, Tempe Mendoan, Tahu Bakso, Bakwan Jagung, serta aneka minuman segar dengan harga yang sangat terjangkau mulai dari Rp 3.000,-.")
add_body("Meskipun memiliki cita rasa nikmat dan basis pelanggan setia, proses operasional pemesanan dan pencatatan transaksi di Kedai Marjuki'S selama ini masih berjalan secara manual (konvensional) menggunakan kertas nota fisik. Kendala yang sering terjadi antara lain:")
add_bullet("Terjadinya antrean penumpukan pesanan saat jam sibuk karena pelanggan harus mencatat manual satu per satu.", "1. Antrean & Waktu Tunggu: ")
add_bullet("Risiko human-error dalam perhitungan nota harga dan keterbatasan rekap keuangan harian yang rentan hilang atau rusak.", "2. Rekapitulasi Manual: ")
add_bullet("Belum tersedianya katalog menu digital berfoto yang memudahkan pelanggan melihat ketersediaan stok porsi secara transparan.", "3. Informasi Menu Terbatas: ")
add_bullet("Tuntutan adaptasi transaksi non-tunai (Cashless) melalui standardisasi QRIS Nasional (Quick Response Code Indonesian Standard) untuk mempercepat pembayaran.", "4. Pembayaran Digital: ")
add_body("Berdasarkan permasalahan tersebut, dirancanglah sebuah sistem informasi pemesanan mandiri (Self-Ordering System) berbasis web modern bertajuk 'Web E-Commerce Kedai Marjuki'S' yang memudahkan pembeli memesan secara Dine-In maupun Takeaway serta membantu pemilik mengelola pesanan secara otomatis.")

add_header_2("1.2 Rumusan Masalah")
add_body("Berdasarkan latar belakang di atas, rumusan masalah dalam pengembangan sistem ini adalah:")
add_bullet("Bagaimana merancang sistem pemesanan kuliner berbasis web responsif yang mudah diakses melalui smartphone pelanggan tanpa memerlukan instalasi aplikasi rumit?")
add_bullet("Bagaimana mengintegrasikan metode pembayaran ganda (QRIS Digital Statis/Dinamis dan Tunai di Meja) secara otomatis dengan verifikasi dua arah?")
add_bullet("Bagaimana membangun dashboard admin terpadu dengan notifikasi audio real-time dan rekapitulasi omset otomatis untuk efisiensi operasional kedai?")

add_header_2("1.3 Maksud dan Tujuan")
add_body("Maksud dari pembuatan proposal dan pengembangan sistem ini adalah mewujudkan transformasi digital UMKM kuliner rumahan menuju sistem operasional yang modern, cepat, dan transparan.")
add_body("Adapun tujuan spesifik proyek ini adalah:")
add_bullet("Menyediakan aplikasi web e-commerce pemesanan mandiri (Self-Ordering) yang mendukung penyajian 'Makan di Tempat' dan 'Bungkus (Bawa Pulang)'.")
add_bullet("Memfasilitasi pembayaran instan menggunakan QRIS GPN (GoPay, OVO, ShopeePay, Dana, m-Banking) dan Uang Tunai.")
add_bullet("Menyediakan fitur live stock tracking dan 1-klik penyelesaian pesanan bagi pemilik kedai.")
add_bullet("Menghasilkan laporan keuangan rekapitulasi omset otomatis berdasarkan periode waktu harian, mingguan, dan bulanan.")

add_header_2("1.4 Manfaat Sistem (Pelanggan, Kasir, Admin, Owner)")
add_bullet("Memperoleh kemudahan melihat katalog 9 menu berfoto, memesan dari meja/rumah, memilih porsi makanan, dan membayar instan via QRIS tanpa antre.", "a. Bagi Pelanggan (Customer): ")
add_bullet("Mendapatkan notifikasi suara lonceng otomatis saat pesanan masuk, memvalidasi bukti bayar dengan cepat, dan menyelesaikan pesanan dalam 1 kali klik.", "b. Bagi Kasir / Pelayan: ")
add_bullet("Kemudahan mengelola katalog makanan (tambah menu, ganti harga, update stok), upload foto QRIS baru, dan memantau database pelanggan.", "c. Bagi Administrator Kedai: ")
add_bullet("Memperoleh transparansi omset harian secara real-time, mengetahui hidangan paling laris (Best Seller), dan meminimalkan kebocoran kas.", "d. Bagi Pemilik Usaha (Owner): ")

add_header_2("1.5 Ruang Lingkup & Batasan Masalah")
add_bullet("Sistem berfokus pada operasional Kedai Marjuki'S Candisari Semarang (Makan di Tempat & Bungkus/Ambil di Kedai).")
add_bullet("Sistem berbasis web responsif (Web Application) yang dioptimasi untuk perangkat Smartphone (Mobile-First) dan Desktop/Laptop.")
add_bullet("Metode pembayaran mencakup QRIS Nasional dan Pembayaran Tunai Langsung.")
add_bullet("Data disimpan menggunakan basis data relasional SQLite berkinerja tinggi.")

doc.add_page_break()

# ----------------- BAB II -----------------
add_header_1("BAB II: LANDASAN TEORI & SPESIFIKASI TEKNIS")

add_header_2("2.1 Landasan Teori Teknologi")
add_bullet("Bahasa pemrograman backend server-side modern yang mendukung paradigma Object-Oriented Programming (OOP), Type Safety, dan eksekusi cepat.", "1. PHP 8.2+: ")
add_bullet("Framework PHP terkemuka dengan arsitektur Model-View-Controller (MVC), Eloquent ORM, Blade Templating Engine, Service Layer Pattern, dan keamanan CSRF/XSS bawaan.", "2. Laravel 11: ")
add_bullet("Sistem manajemen basis data relasional (RDBMS) serverless, zero-configuration, dan self-contained yang andal untuk operasional kedai lokal.", "3. SQLite Database: ")
add_bullet("Framework CSS & Komponen UI modern standar industri yang menjamin tampilan responsif, estetik, dan kompatibel 100% saat diakses offline maupun online.", "4. Bootstrap 5.3: ")
add_bullet("Standardisasi pembayaran digital Bank Indonesia (NMID: ID1026504236773 - LAPAK MARJUKI) yang mendukung seluruh Penyelenggara Jasa Pembayaran (PJP).", "5. QRIS & GPN: ")

add_header_2("2.2 Spesifikasi Kebutuhan Perangkat Keras (Hardware)")
add_bullet("Processor Intel Core i3 / AMD Ryzen 3, RAM 4 GB, Storage SSD 128 GB, Display 1366x768.", "a. Sisi Server / Laptop Admin: ")
add_bullet("Smartphone Android / iOS dengan kamera scan QR, RAM 2 GB+, koneksi internet 4G/Wi-Fi.", "b. Sisi Klien / Smartphone Pelanggan: ")

add_header_2("2.3 Spesifikasi Kebutuhan Perangkat Lunak (Software)")
add_bullet("Sistem Operasi: Windows 10/11 / Linux / macOS.")
add_bullet("Runtime Environment: PHP >= 8.2, Composer 2.x.")
add_bullet("Web Server: Built-in PHP Server / Nginx / Apache.")
add_bullet("Web Browser: Google Chrome, Microsoft Edge, Safari, Mozilla Firefox.")

add_header_2("2.4 Matriks Hak Akses Pengguna (User Roles Matrix)")
roles_table = doc.add_table(rows=1, cols=4)
roles_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = roles_table.rows[0].cells
headers = ["Fitur & Modul Sistem", "Tamu (Guest)", "Pelanggan (Buyer)", "Admin / Owner"]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    set_cell_shading(hdr_cells[i], "B91C1C")
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

matrix_data = [
    ("Melihat Halaman Beranda & Kontak", "Ya", "Ya", "Ya"),
    ("Melihat & Mencari Katalog Menu", "Ya", "Ya", "Ya"),
    ("Menambah Menu ke Keranjang Belanja", "Ya", "Ya", "Ya"),
    ("Checkout & Pilih Makan di Tempat/Bungkus", "Harus Login", "Ya", "Ya"),
    ("Membayar Scan QRIS & Konfirmasi", "Tidak", "Ya", "Ya"),
    ("Melihat Dashboard Omset & Statistik", "Tidak", "Tidak", "Ya"),
    ("Kelola Menu Makanan (CRUD Produk)", "Tidak", "Tidak", "Ya"),
    ("Verifikasi Lunas & Tolak Pembayaran", "Tidak", "Tidak", "Ya"),
    ("1-Klik Selesaikan Pesanan", "Tidak", "Tidak", "Ya"),
    ("Ganti Foto QRIS & Rekap Laporan", "Tidak", "Tidak", "Ya")
]
for row in matrix_data:
    row_cells = roles_table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = val
        p = row_cells[i].paragraphs[0]
        if i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(row_cells[i], 60, 60, 100, 100)

doc.add_page_break()

# ----------------- BAB III -----------------
add_header_1("BAB III: PERANCANGAN SISTEM (SYSTEM DESIGN)")

add_header_2("3.1 Context Diagram (Diagram Konteks Level 0)")
add_body("Diagram Konteks menggambarkan batasan sistem informasi Kedai Marjuki'S yang berinteraksi dengan dua entitas eksternal utama: Pelanggan (Buyer) dan Pengelola (Admin/Kasir).")
add_bullet("Input dari Pelanggan: Registrasi akun, data pesanan (Makan di Tempat/Bungkus), dan konfirmasi bayar QRIS/Tunai.")
add_bullet("Output ke Pelanggan: Katalog menu berfoto, status ketersediaan stok, nomor pesanan, barcode QRIS, dan struk rincian hidangan.")
add_bullet("Input dari Admin: Pengelolaan menu (CRUD), upload foto QRIS, validasi status pembayaran, dan penyesuaian status dapur.")
add_bullet("Output ke Admin: Notifikasi audio pesanan baru, daftar verifikasi transaksi, database pelanggan, dan laporan omset.")

add_header_2("3.2 Data Flow Diagram (DFD Level 1)")
add_body("DFD Level 1 merinci aliran data ke dalam 5 proses fundamental:")
add_bullet("Proses 1.0 (Autentikasi & Registrasi): Mengelola kredensial akun pengguna ke Data Store Users.")
add_bullet("Proses 2.0 (Pengelolaan Menu): Mengelola data hidangan ke Data Store Products & Categories.")
add_bullet("Proses 3.0 (Pemesanan & Keranjang): Memproses item menu terpilih ke Data Store Orders & OrderItems.")
add_bullet("Proses 4.0 (Verifikasi Pembayaran): Memvalidasi transaksi QRIS/Tunai ke Data Store Payments & QRIS.")
add_bullet("Proses 5.0 (Pelaporan & Omset): Menghitung agregasi transaksi lunas menjadi laporan keuangan.")

add_header_2("3.3 Use Case Diagram & Skenario Aktor")
add_body("Sistem memiliki 2 Aktor utama dengan use case sebagai berikut:")
add_bullet("Aktor Pelanggan (Buyer): Mengakses Landing Page, Mencari Menu, Mengelola Keranjang, Checkout Pesanan, Scan QRIS, Melihat Status Masak, dan Mengubah Profil.", "1. ")
add_bullet("Aktor Administrator (Admin): Login Admin, Memantau Dashboard Omset, Menambah/Mengubah/Menghapus Menu, Memverifikasi Pembayaran Lunas, 1-Klik Selesaikan Pesanan, Mengganti Foto QRIS, dan Melihat Rekap Laporan.", "2. ")

add_header_2("3.4 Conceptual Data Model (CDM)")
add_body("CDM memodelkan entitas bisnis utama: USER (Pengguna), CATEGORY (Kategori Menu), PRODUCT (Menu Makanan/Minuman), ORDER (Pesanan), ORDER_ITEM (Detail Rincian Item), dan QRIS_SETTING (Pengaturan QRIS). Hubungan antar entitas mencakup 1-to-N antara User ke Order, 1-to-N antara Category ke Product, dan 1-to-N antara Order ke OrderItem.")

add_header_2("3.5 Physical Data Model (PDM)")
add_body("PDM mendefinisikan tabel fisik database SQLite:")
add_bullet("users (id INT PK, name VARCHAR, email VARCHAR, phone VARCHAR, address TEXT, password VARCHAR, role VARCHAR, created_at TIMESTAMP)")
add_bullet("categories (id INT PK, name VARCHAR, slug VARCHAR, created_at TIMESTAMP)")
add_bullet("products (id INT PK, category_id INT FK, name VARCHAR, slug VARCHAR, description TEXT, price DECIMAL, stock INT, image VARCHAR, status VARCHAR)")
add_bullet("orders (id INT PK, user_id INT FK, order_number VARCHAR, total_price DECIMAL, dining_option VARCHAR, payment_method VARCHAR, payment_status VARCHAR, order_status VARCHAR, notes TEXT)")
add_bullet("order_items (id INT PK, order_id INT FK, product_id INT FK, quantity INT, price DECIMAL, subtotal DECIMAL)")
add_bullet("qris_settings (id INT PK, description VARCHAR, qris_image VARCHAR, is_active BOOLEAN)")

add_header_2("3.6 Entity Relationship Diagram (ERD Notasi Chen)")
add_body("ERD Notasi Chen mengilustrasikan Entitas persegi panjang (USER, PRODUCT, ORDER), Atribut berbentuk oval (id, name, price, total_price), dan Relasi belah ketupat (Melakukan Pemesanan, Memiliki Kategori, Memuat Item) dengan kardinalitas One-to-Many (1:N).")

add_header_2("3.7 Flowchart Workflow Operasional & Shift")
add_body("Alur operasional dimulai dari Pelanggan membuka web ➔ Memilih menu ➔ Checkout (Dine-in / Takeaway) ➔ Scan QRIS / Tunai ➔ Notifikasi lonceng berbunyi di laptop Admin ➔ Dapur memasak ➔ Admin klik '1-Klik Pesanan Selesai' ➔ Makanan disajikan ➔ Selesai.")

doc.add_page_break()

# ----------------- BAB IV -----------------
add_header_1("BAB IV: RENCANA ANGGARAN BIAYA (RAB) & JADWAL PELAKSANAAN")

add_header_2("4.1 Rencana Anggaran Biaya (RAB Pengembangan & Maintenance)")
rab_table = doc.add_table(rows=1, cols=5)
rab_table.alignment = WD_TABLE_ALIGNMENT.CENTER
r_hdr = rab_table.rows[0].cells
r_headers = ["No", "Komponen Biaya / Item", "Volume", "Harga Satuan (Rp)", "Total Biaya (Rp)"]
for i, h in enumerate(r_headers):
    r_hdr[i].text = h
    set_cell_shading(r_hdr[i], "B91C1C")
    p = r_hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

rab_items = [
    ("1", "Analisa Kebutuhan & Desain UI/UX (Figma)", "1 Paket", "500.000", "500.000"),
    ("2", "Pengembangan Web Backend & Database (Laravel 11)", "1 Paket", "1.500.000", "1.500.000"),
    ("3", "Pengembangan Frontend Responsif & Komponen Bootstrap 5.3", "1 Paket", "1.000.000", "1.000.000"),
    ("4", "Integrasi QRIS GPN & Notifikasi Suara Real-Time", "1 Paket", "500.000", "500.000"),
    ("5", "Domain .com / .id & Cloud Hosting (1 Tahun)", "1 Tahun", "650.000", "650.000"),
    ("6", "Testing, Deployment, & Pelatihan Kasir Kedai", "1 Paket", "350.000", "350.000"),
    ("7", "Pemeliharaan Sistem & Backup Berkala (Maintenance)", "1 Tahun", "500.000", "500.000"),
    ("", "TOTAL ANGGARAN PENGEMBANGAN", "", "", "5.000.000")
]
for row in rab_items:
    r_cells = rab_table.add_row().cells
    for i, val in enumerate(row):
        r_cells[i].text = val
        p = r_cells[i].paragraphs[0]
        if i in [0, 2]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif i in [3, 4]:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if row[0] == "":
            set_cell_shading(r_cells[i], "F1F5F9")
            if len(p.runs) > 0:
                p.runs[0].font.bold = True
        set_cell_margins(r_cells[i], 60, 60, 100, 100)

add_header_2("4.2 Timeline Pelaksanaan Proyek (Gantt Chart 8 Minggu)")
gantt_table = doc.add_table(rows=1, cols=9)
gantt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
g_hdr = gantt_table.rows[0].cells
g_headers = ["Aktivitas / Minggu", "M1", "M2", "M3", "M4", "M5", "M6", "M7", "M8"]
for i, h in enumerate(g_headers):
    g_hdr[i].text = h
    set_cell_shading(g_hdr[i], "B91C1C")
    p = g_hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

gantt_rows = [
    ("1. Observasi & Wawancara UMKM", "V", "V", "", "", "", "", "", ""),
    ("2. Perancangan Diagram (DFD/ERD/Figma)", "", "V", "V", "", "", "", "", ""),
    ("3. Setup Database & Skema Migrasi", "", "", "V", "V", "", "", "", ""),
    ("4. Koding Backend & Service Layer", "", "", "", "V", "V", "", "", ""),
    ("5. Koding UI Frontend & Mobile View", "", "", "", "", "V", "V", "", ""),
    ("6. Integrasi QRIS & Notifikasi Suara", "", "", "", "", "", "V", "V", ""),
    ("7. Black Box Testing & Uji Coba", "", "", "", "", "", "", "V", "V"),
    ("8. Deployment & Serah Terima Sistem", "", "", "", "", "", "", "", "V")
]
for row in gantt_rows:
    g_cells = gantt_table.add_row().cells
    for i, val in enumerate(row):
        g_cells[i].text = val
        p = g_cells[i].paragraphs[0]
        if i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if val == "V":
                set_cell_shading(g_cells[i], "DC2626")
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                p.runs[0].font.bold = True
        set_cell_margins(g_cells[i], 50, 50, 80, 80)

doc.add_page_break()

# ----------------- BAB V -----------------
add_header_1("BAB V: PENGUJIAN SISTEM (BLACK BOX TESTING)")

add_header_2("5.1 Metode Pengujian Sistem")
add_body("Pengujian sistem dilakukan menggunakan metode Black Box Testing (Pengujian Fungsionalitas). Metode ini berfokus pada pengujian input dan output tanpa harus melihat struktur kode internal, guna memastikan seluruh fungsionalitas tombol, form, validasi, dan alur transaksi bekerja 100% sesuai spesifikasi kebutuhan pengguna.")

add_header_2("5.2 Matriks Skenario Test Case Black Box")
test_table = doc.add_table(rows=1, cols=5)
test_table.alignment = WD_TABLE_ALIGNMENT.CENTER
t_hdr = test_table.rows[0].cells
t_headers = ["ID", "Skenario Pengujian", "Langkah Aksi (Input)", "Hasil yang Diharapkan (Output)", "Status"]
for i, h in enumerate(t_headers):
    t_hdr[i].text = h
    set_cell_shading(t_hdr[i], "B91C1C")
    p = t_hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

test_cases = [
    ("TC-01", "Pencarian Menu Makanan", "Ketik 'Soto' pada Search Bar katalog", "Menampilkan hanya menu Soto Ayam", "BERHASIL (Valid)"),
    ("TC-02", "Filter Kategori Menu", "Klik pil filter 'Minuman'", "Menampilkan menu Es Teh & Es Jeruk", "BERHASIL (Valid)"),
    ("TC-03", "Tambah ke Keranjang", "Pilih porsi 2 ➔ Klik + Keranjang", "Keranjang bertambah 2 item & subtotal pas", "BERHASIL (Valid)"),
    ("TC-04", "Checkout Makan di Tempat", "Isi Nama, No WA, pilih Makan di Tempat", "Pesanan terbentuk dengan status Menunggu", "BERHASIL (Valid)"),
    ("TC-05", "Pembayaran QRIS", "Klik 'Saya Sudah Membayar' pada QRIS", "Status berubah jadi 'waiting_confirmation'", "BERHASIL (Valid)"),
    ("TC-06", "Notifikasi Suara Masuk", "Pelanggan submit pesanan baru", "Laptop admin berbunyi lonceng instan", "BERHASIL (Valid)"),
    ("TC-07", "Verifikasi Pembayaran Admin", "Admin klik 'Konfirmasi LUNAS'", "Status bayar Lunas & pesanan otomatis Selesai", "BERHASIL (Valid)"),
    ("TC-08", "Pengurangan Stok Otomatis", "Pesanan 2 porsi dikonfirmasi lunas", "Stok menu di database berkurang 2 porsi", "BERHASIL (Valid)"),
    ("TC-09", "Laporan Rekap Omset", "Filter laporan periode 'Hari Ini'", "Total omset terhitung otomatis sesuai kas", "BERHASIL (Valid)"),
    ("TC-10", "Ganti Foto QRIS Admin", "Upload gambar QRIS LAPAK MARJUKI baru", "Gambar QRIS di sisi pembeli terupdate", "BERHASIL (Valid)")
]
for row in test_cases:
    t_cells = test_table.add_row().cells
    for i, val in enumerate(row):
        t_cells[i].text = val
        p = t_cells[i].paragraphs[0]
        if i in [0, 4]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 4:
                set_cell_shading(t_cells[i], "DCFCE7")
                p.runs[0].font.color.rgb = RGBColor(22, 101, 52)
                p.runs[0].font.bold = True
        set_cell_margins(t_cells[i], 50, 50, 80, 80)

doc.add_page_break()

# ----------------- BAB VI -----------------
add_header_1("BAB VI: PENUTUP")

add_header_2("6.1 Kesimpulan")
add_body("Berdasarkan hasil analisa, perancangan, implementasi, dan pengujian Black Box yang telah dilakukan, dapat disimpulkan bahwa:")
add_bullet("Sistem Informasi E-Commerce & Self-Ordering Kedai Marjuki'S berbasis Laravel 11 berhasil dibangun dan berfungsi 100% sesuai kebutuhan operasional UMKM kuliner rumahan.")
add_bullet("Fitur pemilihan penyajian 'Makan di Tempat' dan 'Bungkus' memudahkan pelanggan memesan makanan tanpa antre dan tanpa kebingungan nomor meja.")
add_bullet("Integrasi pembayaran QRIS LAPAK MARJUKI dan Tunai yang dilengkapi notifikasi suara lonceng otomatis terbukti meningkatkan kecepatan pelayanan dan akurasi kas.")
add_bullet("Penggunaan framework modern Laravel 11 dengan arsitektur MVC dan Service Layer menjamin kode program rapi, modular, aman, dan mudah dikembangkan di masa mendatang.")

add_header_2("6.2 Saran & Rencana Pengembangan Masa Depan")
add_body("Untuk pengembangan sistem lebih lanjut di masa mendatang, disarankan:")
add_bullet("Integrasi Payment Gateway Otomatis (seperti Midtrans / Xendit) untuk verifikasi pembayaran QRIS secara real-time callback API tanpa perlu konfirmasi manual.")
add_bullet("Implementasi integrasi kurir pengiriman instan (seperti GrabExpress / Gosend API) jika kedai ingin membuka layanan pesan antar ke rumah pelanggan.")
add_bullet("Pengembangan fitur WhatsApp Gateway Notifikasi otomatis untuk mengirimkan pesan struk pesanan langsung ke nomor handphone pembeli.")

# Save documents
doc_path = r"C:\Users\user\.gemini\antigravity\scratch\kedai-marjukis\Proposal_Kedai_Marjukis_Lengkap.docx"
artifact_doc_path = r"C:\Users\user\.gemini\antigravity\brain\3421a12b-a7ed-42b7-b7f4-8c5a094bbd00\Proposal_Kedai_Marjukis_Lengkap.docx"

doc.save(doc_path)
doc.save(artifact_doc_path)

print(f"SUCCESS: Saved to {doc_path} and {artifact_doc_path}")
